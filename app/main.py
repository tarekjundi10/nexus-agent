import os
from dotenv import load_dotenv
load_dotenv()

import uuid
from fastapi import FastAPI
from fastapi.responses import HTMLResponse
from pydantic import BaseModel
from fastapi.background import BackgroundTasks
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Pt, RGBColor
import io

load_dotenv()

app = FastAPI()

sessions = {}


class ResearchRequest(BaseModel):
    goal: str


def run_agent(session_id: str, goal: str):
    try:
        from agent.graph import build_graph
        sessions[session_id]["status"] = "running"
        sessions[session_id]["steps"] = []

        agent = build_graph(
            step_callback=lambda step: sessions[session_id]["steps"].append(step)
        )

        result = agent.invoke({
            "goal": goal,
            "queries": [],
            "findings": [],
            "report": "",
            "steps": [],
            "reflection": "",
            "iteration": 0
        })

        sessions[session_id]["status"] = "complete"
        sessions[session_id]["report"] = result["report"]

    except Exception as e:
        sessions[session_id]["status"] = "error"
        sessions[session_id]["error"] = str(e)


@app.post("/research")
def start_research(request: ResearchRequest, background_tasks: BackgroundTasks):
    session_id = str(uuid.uuid4())
    sessions[session_id] = {
        "goal": request.goal,
        "status": "queued",
        "steps": [],
        "report": "",
        "error": ""
    }
    background_tasks.add_task(run_agent, session_id, request.goal)
    return {"session_id": session_id}


@app.get("/status/{session_id}")
def get_status(session_id: str):
    if session_id not in sessions:
        return {"error": "Session not found"}
    s = sessions[session_id]
    return {
        "status": s["status"],
        "steps": s["steps"],
        "error": s.get("error", "")
    }


@app.get("/report/{session_id}")
def get_report(session_id: str):
    if session_id not in sessions:
        return {"error": "Session not found"}
    s = sessions[session_id]
    return {"status": s["status"], "goal": s["goal"], "report": s["report"]}


@app.get("/sessions")
def get_sessions():
    result = []
    for sid, s in sessions.items():
        result.append({
            "session_id": sid,
            "goal": s["goal"],
            "status": s["status"]
        })
    return result

@app.get("/download/{session_id}")
def download_docx(session_id: str):
    if session_id not in sessions:
        return {"error": "Session not found"}
    
    s = sessions[session_id]
    report = s.get("report", "")
    goal = s.get("goal", "Research Report")

    doc = Document()

    title = doc.add_heading(goal, 0)
    title.runs[0].font.color.rgb = RGBColor(0x11, 0x11, 0x18)

    for line in report.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            import re
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line[2:])
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')
        
        elif line.startswith('---'):
            doc.add_paragraph('─' * 50)
        else:
            import re
            line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            line = re.sub(r'\*(.*?)\*', r'\1', line)
            if line.strip():
                doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = goal[:40].replace(' ', '_') + '.docx'
    return StreamingResponse(
        buffer,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'}
    )
    
@app.get("/", response_class=HTMLResponse)
def index():
    with open("app/index.html", "r") as f:
        return f.read()